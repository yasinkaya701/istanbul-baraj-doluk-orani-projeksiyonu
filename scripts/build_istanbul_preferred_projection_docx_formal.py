#!/usr/bin/env python3
from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path('/Users/yasinkaya/Hackhaton')
PREFERRED_DIR = ROOT / 'output' / 'istanbul_preferred_projection_2040'
RISK_DIR = PREFERRED_DIR / 'operational_risk_2030'
BENCH_PATH = ROOT / 'output' / 'istanbul_hybrid_physics_sourceaware_ensemble_2040' / 'ensemble_vs_benchmark_models.csv'
ITER_PATH = ROOT / 'output' / 'model_iteration_comparison_2026_03_12.csv'
BOTTOM_UP_SUMMARY_PATH = ROOT / 'output' / 'nearterm_bottom_up_reconciliation' / 'nearterm_bottom_up_reconciliation_summary.json'
SOURCE_REGISTRY = ROOT / 'research' / 'baraj_doluluk_hub' / 'registry' / 'sources' / 'external_sources.csv'
OUT_DOCX = ROOT / 'output' / 'doc' / 'istanbul_baraj_tercih_edilen_projeksiyon_raporu_resmi.docx'
USED_SOURCE_IDS = ['SRC-001', 'SRC-002', 'SRC-057', 'SRC-059', 'SRC-064', 'SRC-065', 'SRC-066', 'SRC-067', 'SRC-068', 'SRC-072', 'SRC-073']
SCENARIO_LABELS = {
    'wet_mild': 'Ilık-ıslak',
    'management_improvement': 'Yönetim iyileşme',
    'base': 'Temel',
    'hot_dry_high_demand': 'Sıcak-kurak-yüksek talep',
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def style_table(table, header_fill: str = 'D9EAF7') -> None:
    table.style = 'Table Grid'
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_fill)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True


def add_caption(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)


def add_para(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style='List Bullet')
    p.add_run(text)


def fmt_pct(v, digits: int = 2) -> str:
    if v is None or v == '' or pd.isna(v):
        return '-'
    return f'%{float(v):.{digits}f}'


def fmt_num(v, digits: int = 2) -> str:
    if v is None or v == '' or pd.isna(v):
        return '-'
    return f'{float(v):.{digits}f}'


def load_data() -> dict[str, object]:
    det = pd.read_csv(PREFERRED_DIR / 'deterministic_summary.csv')
    prob = pd.read_csv(PREFERRED_DIR / 'probabilistic_summary.csv')
    compare = pd.read_csv(PREFERRED_DIR / 'compare_with_baseprob.csv')
    nearterm = pd.read_csv(RISK_DIR / 'preferred_nearterm_risk_summary_2026_2030.csv')
    bench = pd.read_csv(BENCH_PATH)
    iterations = pd.read_csv(ITER_PATH)
    bottom_up = json.loads(BOTTOM_UP_SUMMARY_PATH.read_text(encoding='utf-8'))
    refs = pd.read_csv(SOURCE_REGISTRY)
    refs = refs[refs['source_id'].isin(USED_SOURCE_IDS)].copy()
    refs['sort_key'] = refs['source_id'].str.extract(r'(\d+)').astype(int)
    refs = refs.sort_values('sort_key')
    return {
        'det': det,
        'prob': prob,
        'compare': compare,
        'nearterm': nearterm,
        'bench': bench,
        'iterations': iterations,
        'bottom_up': bottom_up,
        'refs': refs,
    }


def main() -> None:
    data = load_data()
    det = data['det']
    prob = data['prob']
    compare = data['compare']
    nearterm = data['nearterm']
    bench = data['bench']
    iterations = data['iterations']
    bottom_up = data['bottom_up']
    refs = data['refs']

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.72)
    sec.right_margin = Inches(0.72)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('İstanbul Toplam Baraj Doluluğu İçin Tercih Edilen Projeksiyon Yaklaşımı')
    r.bold = True
    r.font.size = Pt(17.5)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Yöntemsel tercih, test sonuçları ve kullanım çerçevesi')
    r2.italic = True
    r2.font.size = Pt(10.5)

    doc.add_heading('Yönetici özeti', level=1)
    add_para(doc, 'Bu metin, İstanbul toplam baraj doluluğu için geliştirilen projeksiyon sisteminde neden belirli bir model ailesinin tercih edildiğini, bu tercihin hangi test bulgularına dayandığını ve çıktının hangi sınırlar içinde yorumlanması gerektiğini resmî bir çerçevede özetlemektedir. Tercih edilen çözüm, yakın vadede baraj bazlı tahminlerle uzlaştırılmış bir yol ile orta ve uzun vadede fizik-kısıtlı ve kaynak duyarlı hibrit ensemble omurgasını birlikte kullanmaktadır. Böylece hem kısa dönem operasyonel gerçekçilik hem de daha uzun vadeli yapısal tutarlılık korunmaya çalışılmıştır.')

    chosen = bench[bench['model'] == 'hybrid_physics_ensemble_phys'].iloc[0]
    t = doc.add_table(rows=2, cols=3)
    style_table(t, '0F766E')
    for i, h in enumerate(['Tek adım hata', 'Çok adımlı ortalama hata', 'Fiziksel işaret testi']):
        t.cell(0, i).text = h
    vals = [fmt_num(chosen['one_step_rmse_pp']), fmt_num(chosen['mean_recursive_rmse_pp']), '4/4']
    for i, v in enumerate(vals):
        t.cell(1, i).text = v
        set_cell_shading(t.cell(1, i), 'CCFBF1')

    doc.add_heading('1. Tercih edilen yöntem', level=1)
    add_para(doc, 'Projeksiyon çerçevesi, yalnızca geçmiş doluluk serisini ileri taşıyan istatistiksel bir tahmin yaklaşımı olarak tasarlanmamıştır. Bunun yerine sistem; yakın vade için kapasite ağırlıklı bottom-up tahminlerle uzlaştırılmış bir yol, orta ve uzun vade için ise fizik-kısıtlı hibrit ensemble omurgası üzerinden kurulmuştur. Bu tercih, hata minimizasyonunun yanında fiziksel tutarlılığı da zorunlu bir kriter olarak kabul etmektedir.')
    add_bullet(doc, 'Ana model: `hybrid_physics_ensemble_phys`.')
    add_bullet(doc, 'Yakın vade düzeltmesi: baraj bazlı tahminlerin kapasite ağırlıklı toplamı ile uzlaştırma.')
    add_bullet(doc, 'Çıktı tipi: deterministik yol + olasılıksal bant + eşik altı risk analizi.')

    doc.add_heading('2. Model seçimi ve gerekçe', level=1)
    add_para(doc, 'Model seçimi yapılırken tek adımlı başarı ile çok adımlı kararlılık birlikte değerlendirilmiştir. Ayrıca yağış, talep, referans evapotranspirasyon ve transfer değişikliklerine verilen tepkinin fiziksel işaretleri de kontrol edilmiştir. Bu değerlendirme sonucunda `hybrid_physics_ensemble_phys` hem kabul edilebilir hata düzeyi hem de tam fiziksel işaret başarısı nedeniyle tercih edilmiştir.')
    table = doc.add_table(rows=1, cols=4)
    style_table(table)
    for i, h in enumerate(['Model', 'Tek adım hata', 'Çok adımlı hata', 'Fizik testi']):
        table.cell(0, i).text = h
    for row in bench.sort_values('one_step_rmse_pp').head(6).itertuples(index=False):
        r = table.add_row().cells
        r[0].text = row.model
        r[1].text = fmt_num(row.one_step_rmse_pp)
        r[2].text = fmt_num(row.mean_recursive_rmse_pp)
        r[3].text = f"{int(row.physics_pass_count)}/4"
    v5 = iterations[iterations['model'] == 'ensemble_v5_sourcerain_phys'].iloc[0]
    add_para(doc, f"Kaynak bazlı yağış forcing içeren `ensemble_v5_sourcerain_phys` kısa ufukta {v5['one_step_rmse_pp']:.2f} ile daha düşük hata vermesine rağmen, çok adımlı ortalama hata {v5['mean_recursive_rmse_pp']:.2f} seviyesine yükselmiştir. Bu nedenle söz konusu sürüm, araştırma amaçlı bir dal olarak bırakılmış ve ana projeksiyon hattına dahil edilmemiştir. Bu tercih, tek bir metrikte elde edilen sınırlı iyileşmenin genel model kararlılığını bozmasına izin verilmemesi ilkesine dayanmaktadır.")

    doc.add_heading('3. Yakın vade uzlaştırması', level=1)
    add_para(doc, 'Top-down temel yol ile baraj bazlı bottom-up tahminler karşılaştırıldığında, yakın vadede sistematik bir yukarı yönlü sapma görülmüştür. Bu fark doğrudan raporlanmış ve 2026-2029 dönemine uygulanan, zamanla azalan bir düzeltme eğrisi biçiminde ana yola işlenmiştir. Böylece kısa dönem operasyonel görünüm daha temkinli bir seviyeye çekilmiş, 2030 sonrası ise ana fizik-kısıtlı omurga korunmuştur.')
    add_bullet(doc, f"Bottom-up ile top-down temel yol arasındaki ortalama fark: {bottom_up['mean_gap_bottom_up_minus_topdown_pp']:.2f} yüzde puan.")
    add_bullet(doc, f"Gözlenen en sert yakın vade farkı: {bottom_up['min_gap_bottom_up_minus_topdown_pp']:.2f} yüzde puan.")
    add_bullet(doc, 'Düzeltme tasarımı gereği zamanla sönümlenmekte ve uzun vade omurgasını değiştirmemektedir.')
    for path, caption in [
        (ROOT / 'output' / 'nearterm_bottom_up_reconciliation' / 'nearterm_bottom_up_vs_topdown.png', 'Şekil 1. Kapasite ağırlıklı bottom-up tahminler ile top-down temel yol arasındaki fark.'),
        (PREFERRED_DIR / 'reconciled_base_projection.png', 'Şekil 2. Uzlaştırılmış temel senaryo yolu.'),
    ]:
        doc.add_picture(str(path), width=Inches(6.0))
        add_caption(doc, caption)

    doc.add_heading('4. Senaryo sonuçları', level=1)
    merged = det.merge(prob, on='scenario', how='left')
    scen = doc.add_table(rows=1, cols=6)
    style_table(scen)
    for i, h in enumerate(['Senaryo', '2040 sonu', '2040 P10', '2040 P50', '2040 P90', 'İlk %40>=50 yılı']):
        scen.cell(0, i).text = h
    for row in merged.itertuples(index=False):
        r = scen.add_row().cells
        r[0].text = SCENARIO_LABELS.get(row.scenario, row.scenario)
        r[1].text = fmt_pct(row.end_fill_2040_12_pct)
        r[2].text = fmt_pct(row.p10_endpoint_2040_12_pct)
        r[3].text = fmt_pct(row.p50_endpoint_2040_12_pct)
        r[4].text = fmt_pct(row.p90_endpoint_2040_12_pct)
        r[5].text = '-' if pd.isna(row.first_year_prob_below_40_ge_50pct) else str(int(row.first_year_prob_below_40_ge_50pct))
    add_para(doc, 'Senaryo sonuçları birlikte okunduğunda, temel senaryoda 2040 medyan doluluk düzeyinin yaklaşık %34.69 seviyesinde kaldığı; yönetim iyileşmesi senaryosunda ise aynı iklim koşulları altında daha yüksek bir doluluk patikasının mümkün olduğu görülmektedir. Buna karşılık sıcak-kurak-yüksek talep senaryosu, sistemin erken dönemde dahi belirgin bir baskı altına girdiği tek keskin bozulma senaryosudur.')
    for path, caption in [
        (PREFERRED_DIR / 'reconciled_probabilistic_fan_paths_2026_2040.png', 'Şekil 3. Olasılıksal senaryo yolları.'),
        (PREFERRED_DIR / 'reconciled_vs_baseprob_base.png', 'Şekil 4. Rekonsiliasyon öncesi ve sonrası temel senaryo medyan yolu.'),
    ]:
        doc.add_picture(str(path), width=Inches(6.1))
        add_caption(doc, caption)

    doc.add_heading('5. 2026-2030 operasyonel risk', level=1)
    nt = doc.add_table(rows=1, cols=6)
    style_table(nt)
    for i, h in enumerate(['Senaryo', 'Tepe %40 altı risk', 'Tepe ay', 'Tepe %30 altı risk', 'Tepe ay', 'En zayıf yaz P50']):
        nt.cell(0, i).text = h
    for row in nearterm.itertuples(index=False):
        r = nt.add_row().cells
        r[0].text = SCENARIO_LABELS.get(row.scenario, row.scenario)
        r[1].text = fmt_pct(row.max_monthly_prob_below_40_2026_2030_pct)
        r[2].text = str(row.peak_month_prob_below_40_date)
        r[3].text = fmt_pct(row.max_monthly_prob_below_30_2026_2030_pct)
        r[4].text = str(row.peak_month_prob_below_30_date)
        r[5].text = f"{int(row.lowest_p50_summer_year_2026_2030)} / {fmt_pct(row.lowest_p50_summer_fill_pct_2026_2030)}"
    add_para(doc, 'Yakın vade risk analizinde temel ve yönetim iyileşme senaryolarının 2030 sonuna kadar kalıcı `%40 altı` stres üretmediği görülmektedir. Buna karşılık sıcak-kurak-yüksek talep senaryosu, 2026 yılından itibaren yoğun risk kümesi üretmektedir. Bu bulgu, kısa dönem alarm seviyesinin her senaryoda aynı olmadığını; özellikle talep ve iklim stresinin aynı anda arttığı koşulların ayrı bir risk sınıfı olarak ele alınması gerektiğini göstermektedir.')
    for path, caption in [
        (RISK_DIR / 'preferred_monthly_threshold_risk_2026_2030.png', 'Şekil 5. 2026-2030 döneminde aylık eşik altı olasılık yolları.'),
        (RISK_DIR / 'preferred_threshold_heatmap_40_2026_2030.png', 'Şekil 6. 2026-2030 aylık %40 altı risk ısı haritası.'),
    ]:
        doc.add_picture(str(path), width=Inches(6.0))
        add_caption(doc, caption)

    doc.add_heading('6. Kullanım sınırı ve sonraki çalışma başlıkları', level=1)
    add_bullet(doc, 'Referans evapotranspirasyon (ET0) bloğu hâlen vekil radyasyon serileriyle çalışmaktadır; aktinograf verisi geldiğinde yeniden kurulması gerekmektedir.')
    add_bullet(doc, 'Açık su yüzeyi buharlaşması ET0 ile özdeş kabul edilmemekte, ancak alan eğrileri ve daha ayrıntılı açık su kayıp bilgisi henüz eksiktir.')
    add_bullet(doc, 'Uzaysal yağış forcing denenmiş, ancak gelecek için güvenilir uzaysal forcing seti olmadan kararlılık bozulmuştur.')
    add_bullet(doc, 'Bu nedenle mevcut paketin en uygun kullanımı, yakın vadede operasyonel okuma ve orta-uzun vadede senaryo karşılaştırması üretmektir; nihai işletme modeli olarak değerlendirilmemelidir.')

    diff = doc.add_table(rows=1, cols=2)
    style_table(diff, 'DCFCE7')
    diff.cell(0,0).text = 'Senaryo'
    diff.cell(0,1).text = 'Rekonsiliasyonun 2040 P50 etkisi (yp)'
    for row in compare[['scenario', 'delta_p50_endpoint_pp']].itertuples(index=False):
        r = diff.add_row().cells
        r[0].text = SCENARIO_LABELS.get(row.scenario, row.scenario)
        r[1].text = fmt_num(row.delta_p50_endpoint_pp)
    add_para(doc, 'Bu tablo, rekonsiliasyon etkisinin 2040 sonu üzerinde sınırlı kaldığını göstermektedir. Bu sonuç beklenen bir durumdur; çünkü düzeltme yakın vade operasyonel görünümü iyileştirmek için tasarlanmış ve uzun vadede sönümlenecek şekilde kurulmuştur.')

    doc.add_heading('Kaynakça', level=1)
    for row in refs.itertuples(index=False):
        doc.add_paragraph(f"[{row.source_id}] {row.title}. {row.organization_or_journal}. {row.year}. {row.url}")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == '__main__':
    main()
