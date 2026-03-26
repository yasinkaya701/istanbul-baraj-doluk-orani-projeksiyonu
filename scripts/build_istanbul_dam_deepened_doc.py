#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Append deepened research section to Istanbul dam DOCX.")
    parser.add_argument(
        "--base-docx",
        type=Path,
        default=Path("output/doc/istanbul_baraj_durum_ozeti_akademik.docx"),
    )
    parser.add_argument(
        "--deep-root",
        type=Path,
        default=Path("output/istanbul_dam_deep_research"),
    )
    parser.add_argument(
        "--out-docx",
        type=Path,
        default=Path("output/doc/istanbul_baraj_durum_ozeti_akademik_derinlesmis.docx"),
    )
    return parser.parse_args()


def add_line(document: Document, text: str, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)


def find_delta(df: pd.DataFrame, scenario: str, horizon: int) -> float:
    value = df.loc[
        (df["scenario"] == scenario) & (df["horizon_month"] == horizon),
        "delta_pp",
    ]
    return float(value.iloc[0])


def add_scenario_table(document: Document, scenario_df: pd.DataFrame) -> None:
    rows = [
        ("+%10 yağış, 3 ay", "rain_plus10_3m"),
        ("+%10 ET0, 3 ay", "et0_plus10_3m"),
        ("+%10 tüketim, 3 ay", "cons_plus10_3m"),
        ("-%7 talep kısıtı, 3 ay", "restriction_minus7_3m"),
        ("-%15 talep kısıtı, 3 ay", "restriction_minus15_3m"),
        ("-%22 talep kısıtı, 3 ay", "restriction_minus22_3m"),
        ("-%15 kısıt + geri tepme", "restriction_minus15_with_rebound"),
        ("Sıcak-kurak-yüksek talep", "hot_dry_high_demand"),
    ]
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Senaryo"
    header[1].text = "1. ay"
    header[2].text = "3. ay"
    header[3].text = "6. ay"
    header[4].text = "12. ay"
    for label, key in rows:
        row = table.add_row().cells
        row[0].text = label
        for idx, horizon in enumerate([1, 3, 6, 12], start=1):
            row[idx].text = f"{find_delta(scenario_df, key, horizon):+.2f} yp"


def add_sources(document: Document, sources: list[dict[str, str]]) -> None:
    for item in sources:
        add_line(document, f"- {item['topic']}: {item['url']}")


def main() -> None:
    args = parse_args()
    deep_root = args.deep_root.resolve()
    fig_dir = deep_root / "figures"
    out_docx = args.out_docx.resolve()
    shutil.copyfile(args.base_docx.resolve(), out_docx)

    metrics = pd.read_csv(deep_root / "model_ablation_metrics.csv")
    scenario_df = pd.read_csv(deep_root / "scenario_summary.csv")
    weighted_df = pd.read_csv(deep_root / "weighted_total_vs_mean.csv")
    sources = json.loads((deep_root / "sources.json").read_text(encoding="utf-8"))
    summary = json.loads((deep_root / "summary.json").read_text(encoding="utf-8"))

    full_rmse = float(metrics.loc[metrics["model"] == "full_model", "rmse_pp"].iloc[0])
    pattern_rmse = float(metrics.loc[metrics["model"] == "pattern_only", "rmse_pp"].iloc[0])
    climate_rmse = float(metrics.loc[metrics["model"] == "climate_plus_memory", "rmse_pp"].iloc[0])
    demand_rmse = float(metrics.loc[metrics["model"] == "demand_plus_memory", "rmse_pp"].iloc[0])

    avg_diff = float(weighted_df["diff_pp"].mean())
    q25 = float(weighted_df["diff_pp"].quantile(0.25))
    q75 = float(weighted_df["diff_pp"].quantile(0.75))

    document = Document(out_docx)
    document.add_page_break()
    document.add_heading("Derinleştirilmiş Araştırma ve Simülasyon Eki", level=1)

    add_line(
        document,
        "Bu ek bölümde iki metodolojik iyileştirme yapıldı: "
        "i) hedef seri eşit ağırlıklı ortalama yerine İSKİ aktif hacimleriyle ağırlıklandırılmış toplam doluluk oranına çevrildi, "
        "ii) aylık seviye yerine aylık değişim modeli kurularak yağış, ET0 ve tüketim etkileri mevsimsellikten arındırılmış 2 aylık anomalilerle yeniden simüle edildi.",
    )
    add_line(
        document,
        "Yağış tarafında new data doğrudan 2011-01 ile 2021-12 arasında var. "
        "2022-01 ile 2024-02 arasında model sürekliliği için eski aylık yağış serisi sadece fallback olarak tutuldu. "
        "Bu ayrım özellikle akademik sunumda açıkça belirtilmelidir.",
    )

    document.add_heading("1. Toplam Doluluk Serisinin Düzeltilmesi", level=2)
    add_line(
        document,
        f"Hacim ağırlıklı toplam doluluk ile eşit ağırlıklı ortalama arasında ortalama fark {avg_diff:.2f} yüzde puan çıktı. "
        f"Çeyrekler arası aralık yaklaşık {q25:.2f} ile {q75:.2f} yüzde puan arasında. "
        "Bu nedenle küçük barajlarla büyük depolamaları aynı etkide gören overall_mean serisi, şehir ölçeği için tek başına yeterli değil.",
    )
    document.add_picture(str(fig_dir / "weighted_total_vs_mean.png"), width=Inches(6.5))
    add_line(document, "Şekil A1. Eşit ağırlıklı ortalama ile hacim ağırlıklı toplam doluluk serisinin karşılaştırması.")

    document.add_heading("2. Ablasyon ve Model Seçimi", level=2)
    add_line(
        document,
        f"Yürüyen testte tam model RMSE={full_rmse:.2f} yüzde puan, pattern-only model RMSE={pattern_rmse:.2f} yüzde puan verdi. "
        f"İyileşme {pattern_rmse - full_rmse:.2f} yüzde puan. "
        f"İklim+hafıza modeli RMSE={climate_rmse:.2f}; talep+hafıza modeli RMSE={demand_rmse:.2f}.",
    )
    add_line(
        document,
        "Yorum: yağış ve ET0 gibi hidro-meteorolojik sürücüler, şehir geneline agregasyon yapılmış tüketim serisine göre daha net bir ek sinyal taşıyor. "
        "Tüketim tek başına zayıf görünüyor; bunun ana nedeni konut-sanayi-sulama ayrımının eldeki seride karışmış olması.",
    )
    document.add_picture(str(fig_dir / "ablation_rmse_weighted.png"), width=Inches(6.2))
    add_line(document, "Şekil A2. Hacim ağırlıklı toplam doluluk için pattern-only ve exogenous model abasyonu.")

    document.add_heading("3. Şok ve Politika Senaryoları", level=2)
    add_line(
        document,
        "Aşağıdaki senaryolar 12 aylık ufukta recursive şekilde simüle edildi. "
        "Baz yol, aylık klimatoloji; şok senaryoları ise ilk 3 ayda yağış, ET0 veya talep üzerinde değişiklik uyguluyor. "
        "Talep kısıtı büyüklükleri literatürde raporlanan gönüllü ve zorunlu kısıt aralıklarına göre seçildi.",
    )
    add_scenario_table(document, scenario_df)
    add_line(
        document,
        f"+%10 yağışın etkisi 3. ayda {find_delta(scenario_df, 'rain_plus10_3m', 3):+.2f} yüzde puan; "
        f"-%15 talep kısıtının etkisi 3. ayda {find_delta(scenario_df, 'restriction_minus15_3m', 3):+.2f} yüzde puan; "
        f"sıcak-kurak-yüksek talep kombinasyonunun etkisi 6. ayda {find_delta(scenario_df, 'hot_dry_high_demand', 6):+.2f} yüzde puan olarak bulundu.",
    )
    add_line(
        document,
        f"Geri tepme senaryosunda, 3 aylık -%15 kısıt sonrasında iki aylık +%5 toparlanma eklendiğinde bile 12. ay etkisi {find_delta(scenario_df, 'restriction_minus15_with_rebound', 12):+.2f} yüzde puan düzeyinde pozitif kalıyor. "
        "Bu sonuç, kısıtların tamamen günlük değil; birkaç aylık işletme penceresinde kalıcı seviye farkı yaratabildiğini gösteriyor.",
    )
    document.add_picture(str(fig_dir / "scenario_paths_weighted.png"), width=Inches(6.5))
    add_line(document, "Şekil A3. 12 aylık senaryo yolları: yağış şoku, talep kısıtı, geri tepme ve sıcak-kurak birleşik stres.")

    document.add_heading("4. Parametre Etkilerinin Yorumu", level=2)
    add_line(
        document,
        "Grup bazlı katsayı büyüklükleri, bu veri çözünürlüğünde en güçlü yapının depo hafızası olduğunu; "
        "buna ek olarak yağış-kuraklık bilgisinin ET0 ve toplam tüketimden daha net bir ek sinyal taşıdığını gösteriyor. "
        "Bu, ET0 önemsiz demek değildir; yalnızca aylık şehir ortalamasında doğrudan gözlenen açık su buharlaşması ve sektör bazlı çekiş serileri olmadığı için ET0 etkisinin kısmen örtük kaldığını gösterir.",
    )
    document.add_picture(str(fig_dir / "group_importance.png"), width=Inches(5.6))
    add_line(document, "Şekil A4. Tam modelde grup bazlı standartlaştırılmış katsayı büyüklükleri.")

    document.add_heading("5. Literatürden Modele Eklenmesi Gereken Parametreler", level=2)
    add_line(
        document,
        "- Açık su yüzeyi alanı ve seviye-alan-eğrisi: rezervuar evaporasyonu için ET0 tek başına yeterli değildir; yüzey alanı değişimi gerekir."
    )
    add_line(
        document,
        "- Gerçek işletme olayları: su kesintisi tarihleri, basınç düşürme, zorunlu kısıtlar ve sonrasındaki geri tepme talep şoklarını etiketlemek için gerekir."
    )
    add_line(
        document,
        "- Sektör ayrımı: konut, sanayi, ticari kullanım ve sulama ayrı seriler halinde girildiğinde tüketim bloğunun açıklayıcılığı ciddi biçimde artar."
    )
    add_line(
        document,
        "- Havzalar arası transfer ve arıtma tesisi-bağlantı matrisi: hangi barajın hangi yakayı ve hangi arıtma tesisini beslediği, şokların mekânsal yayılımını simüle etmek için gerekir."
    )
    add_line(
        document,
        "- Kayıp/kaçak ve illegal çekiş: özellikle tüketim serisini gerçek çekişe yaklaştırmak için ayrıca modellenmelidir."
    )
    add_line(
        document,
        "- Tarımsal çekiş proxy'si: ET0 ile birlikte ürün deseninden türetilen sulama ihtiyacı serisi yaz baskısını daha doğru yakalar."
    )

    document.add_heading("6. Literatür Özeti ve Kaynaklar", level=2)
    add_line(
        document,
        "FAO-56 ve HEC-HMS dokümantasyonu, günlük ölçekte Penman-Monteith kurulumunun temel referansını veriyor. "
        "HESS 2024 ve 2026 çalışmaları, açık su evaporasyonu ile iklim sinyalinin rezervuar bulunurluğunu anlamlı biçimde değiştirebildiğini gösteriyor. "
        "Journal of Hydrology 2023 ve 2025 çalışmaları ise hafıza + exogenous sürücüler + operasyonel koşullandırmanın rezervuar tahmininde faydalı olduğunu destekliyor. "
        "Talep kısıtı literatüründe ise AWE 2024, California 2008 ve Los Angeles 2015 çalışmaları gönüllü ve zorunlu kısıtların farklı büyüklüklerde ama ölçülebilir etkiler yarattığını gösteriyor.",
    )
    add_sources(document, sources)

    document.add_heading("7. Kısa Sonuç", level=2)
    add_line(
        document,
        f"Derinleştirilmiş pakette en iyi model {summary['best_model']} oldu ve RMSE {summary['best_rmse_pp']:.2f} yüzde puana indi. "
        "Akademik olarak savunulabilir ana tez şu: İstanbul toplam doluluğu sadece pattern ile değil, hacim ağırlığı + hidro-meteorolojik sürücüler + talep şokları ile birlikte okunmalı; "
        "ancak sektör bazlı çekiş ve gerçek işletme olayları eklenmeden tüketim bloğu halen eksik kalır.",
    )

    document.save(out_docx)
    print(out_docx)


if __name__ == "__main__":
    main()
