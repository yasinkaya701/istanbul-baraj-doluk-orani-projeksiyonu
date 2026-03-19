#!/usr/bin/env python3
from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path('/Users/yasinkaya/Hackhaton')
PREFERRED_DIR = ROOT / 'output' / 'istanbul_preferred_projection_2040'
RISK_DIR = PREFERRED_DIR / 'operational_risk_2030'
BENCH_PATH = ROOT / 'output' / 'istanbul_hybrid_physics_sourceaware_ensemble_2040' / 'ensemble_vs_benchmark_models.csv'
ITER_PATH = ROOT / 'output' / 'model_iteration_comparison_2026_03_12.csv'
BOTTOM_UP_SUMMARY_PATH = ROOT / 'output' / 'nearterm_bottom_up_reconciliation' / 'nearterm_bottom_up_reconciliation_summary.json'
SOURCE_REGISTRY = ROOT / 'research' / 'baraj_doluluk_hub' / 'registry' / 'sources' / 'external_sources.csv'
OUT_DOCX = ROOT / 'output' / 'doc' / 'istanbul_baraj_tercih_edilen_projeksiyon_raporu.docx'
USED_SOURCE_IDS = ['SRC-001', 'SRC-002', 'SRC-057', 'SRC-059', 'SRC-064', 'SRC-065', 'SRC-066', 'SRC-067', 'SRC-068', 'SRC-072', 'SRC-073']

SCENARIO_LABELS = {
    'wet_mild': 'Ilık-ıslak',
    'management_improvement': 'Yönetim iyileşme',
    'base': 'Temel',
    'hot_dry_high_demand': 'Sıcak-kurak-yüksek talep',
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def style_table(table, header_fill: str = 'D9EAF7') -> None:
    table.style = 'Table Grid'
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_shading(cell, header_fill)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style='List Bullet')
    p.add_run(text)


def add_caption(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def load_data() -> dict[str, object]:
    det = pd.read_csv(PREFERRED_DIR / 'deterministic_summary.csv')
    prob = pd.read_csv(PREFERRED_DIR / 'probabilistic_summary.csv')
    compare = pd.read_csv(PREFERRED_DIR / 'compare_with_baseprob.csv')
    nearterm = pd.read_csv(RISK_DIR / 'preferred_nearterm_risk_summary_2026_2030.csv')
    bench = pd.read_csv(BENCH_PATH)
    iterations = pd.read_csv(ITER_PATH)
    bottom_up = json.loads(BOTTOM_UP_SUMMARY_PATH.read_text(encoding='utf-8'))
    refs = pd.read_csv(SOURCE_REGISTRY)
    refs = refs[refs['source_id'].isin(USED_SOURCE_IDS)].copy()
    refs['sort_key'] = refs['source_id'].str.extract(r'(\d+)').astype(int)
    refs = refs.sort_values('sort_key')
    return {
        'det': det,
        'prob': prob,
        'compare': compare,
        'nearterm': nearterm,
        'bench': bench,
        'iterations': iterations,
        'bottom_up': bottom_up,
        'refs': refs,
    }


def fmt_pct(v: float | int | str | None, digits: int = 2) -> str:
    if v is None or v == '' or pd.isna(v):
        return '-'
    return f'%{float(v):.{digits}f}'


def fmt_num(v: float | int | str | None, digits: int = 2) -> str:
    if v is None or v == '' or pd.isna(v):
        return '-'
    return f'{float(v):.{digits}f}'


def main() -> None:
    data = load_data()
    det = data['det']
    prob = data['prob']
    compare = data['compare']
    nearterm = data['nearterm']
    bench = data['bench']
    iterations = data['iterations']
    bottom_up = data['bottom_up']
    refs = data['refs']

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('İstanbul Baraj Doluluğu Tercih Edilen Projeksiyon Raporu')
    r.bold = True
    r.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run('Kısa vadede uzlaştırılmış, uzun vadede fizik-kısıtlı ve kaynak duyarlı hibrit ensemble paketi').italic = True

    doc.add_paragraph(
        'Bu belge, şu anda kullanılmasını önerdiğimiz projeksiyon paketini özetler. ' \
        'Seçilen çözüm, 2026-2029 döneminde baraj bazlı tahminlerle uzlaştırılmıştır; ' \
        '2040 ufkunda ise fizik-kısıtlı ve kaynak duyarlı ensemble omurgasını korur.'
    )

    chosen = bench[bench['model'] == 'hybrid_physics_ensemble_phys'].iloc[0]
    metrics = doc.add_table(rows=2, cols=3)
    style_table(metrics, '0F766E')
    headers = ['Tek adım hata', 'Çok adımlı ortalama hata', 'Fiziksel işaret testi']
    values = [fmt_num(chosen['one_step_rmse_pp']), fmt_num(chosen['mean_recursive_rmse_pp']), '4/4']
    for i in range(3):
        metrics.cell(0, i).text = headers[i]
        metrics.cell(1, i).text = values[i]
        set_cell_shading(metrics.cell(1, i), 'CCFBF1')

    doc.add_heading('1. Seçilen model ve gerekçe', level=1)
    add_bullet(doc, 'Ana omurga `hybrid_physics_ensemble_phys` olarak seçildi.')
    add_bullet(doc, 'Bu model, istatistiksel hata ile fiziksel tutarlılık arasında en dengeli sonucu veriyor.')
    add_bullet(doc, 'Tek adımda 3.91 yüzde puan hata üretirken, fiziksel işaret testlerinin tamamını geçiyor.')
    add_bullet(doc, 'Kısa vade için ayrıca baraj bazlı tahminlerle uzlaştırma uygulanıyor; böylece yakın dönem yol fazla iyimser kalmıyor.')

    table = doc.add_table(rows=1, cols=4)
    style_table(table)
    hdr = table.rows[0].cells
    hdr[0].text = 'Model'
    hdr[1].text = 'Tek adım hata'
    hdr[2].text = 'Çok adımlı hata'
    hdr[3].text = 'Fizik testi'
    for row in bench.sort_values('one_step_rmse_pp').head(6).itertuples(index=False):
        r = table.add_row().cells
        r[0].text = row.model
        r[1].text = fmt_num(row.one_step_rmse_pp)
        r[2].text = fmt_num(row.mean_recursive_rmse_pp)
        r[3].text = f"{int(row.physics_pass_count)}/4"

    v5 = iterations[iterations['model'] == 'ensemble_v5_sourcerain_phys'].iloc[0]
    doc.add_paragraph(
        f"Not: `ensemble_v5_sourcerain_phys` kısa vadede {v5['one_step_rmse_pp']:.2f} ile daha düşük hata verdi; "
        f"ancak çok adımlı hata {v5['mean_recursive_rmse_pp']:.2f} seviyesine çıktığı için ana hatta alınmadı."
    )

    doc.add_heading('2. Yakın vade uzlaştırması', level=1)
    add_bullet(doc, f"Baraj bazlı bottom-up tahminler ile top-down temel yol arasında ortalama {bottom_up['mean_gap_bottom_up_minus_topdown_pp']:.2f} yüzde puan fark bulundu.")
    add_bullet(doc, 'Bu fark 2026-2029 döneminde azalarak uygulanan bir düzeltme eğrisine çevrildi.')
    add_bullet(doc, 'Amaç, yakın vade operasyonel görünümde aşırı iyimserliği kırmak; 2030 sonrası ise ana fizik-kısıtlı omurgayı korumak.')

    for path, caption, width in [
        (ROOT / 'output' / 'nearterm_bottom_up_reconciliation' / 'nearterm_bottom_up_vs_topdown.png', 'Şekil 1. Baraj bazlı bottom-up tahminlerin kapasite ağırlıklı toplamı ile top-down temel yol arasındaki fark.', 6.1),
        (PREFERRED_DIR / 'reconciled_base_projection.png', 'Şekil 2. Uzlaştırılmış temel senaryo yolu.', 6.1),
    ]:
        doc.add_picture(str(path), width=Inches(width))
        add_caption(doc, caption)

    doc.add_heading('3. Deterministik ve olasılıksal senaryo çıktıları', level=1)
    merged = det.merge(prob, on='scenario', how='left')
    t = doc.add_table(rows=1, cols=6)
    style_table(t)
    headers = ['Senaryo', '2040 sonu', '2040 P10', '2040 P50', '2040 P90', 'İlk %40>=50 yılı']
    for i, h in enumerate(headers):
        t.cell(0, i).text = h
    for row in merged.itertuples(index=False):
        r = t.add_row().cells
        r[0].text = SCENARIO_LABELS.get(row.scenario, row.scenario)
        r[1].text = fmt_pct(row.end_fill_2040_12_pct)
        r[2].text = fmt_pct(row.p10_endpoint_2040_12_pct)
        r[3].text = fmt_pct(row.p50_endpoint_2040_12_pct)
        r[4].text = fmt_pct(row.p90_endpoint_2040_12_pct)
        r[5].text = '-' if pd.isna(row.first_year_prob_below_40_ge_50pct) else str(int(row.first_year_prob_below_40_ge_50pct))

    add_bullet(doc, 'Temel senaryoda 2040 medyan doluluk yaklaşık %34.69 seviyesinde kalıyor.')
    add_bullet(doc, 'Yönetim iyileşme senaryosu, iklim aynı kalsa bile sistemi yukarı taşıyor.')
    add_bullet(doc, 'Sıcak-kurak-yüksek talep senaryosu tek agresif kırılma senaryosu olarak öne çıkıyor.')
    add_bullet(doc, 'Rekonsiliasyon sonrası 2040 uç değerleri neredeyse değişmiyor; fark esas olarak yakın vadede oluşuyor.')

    for path, caption, width in [
        (PREFERRED_DIR / 'reconciled_probabilistic_fan_paths_2026_2040.png', 'Şekil 3. Tercih edilen paketin olasılıksal senaryo yolları.', 6.2),
        (PREFERRED_DIR / 'reconciled_vs_baseprob_base.png', 'Şekil 4. Rekonsiliasyon öncesi ve sonrası temel senaryo medyan yolu.', 6.2),
    ]:
        doc.add_picture(str(path), width=Inches(width))
        add_caption(doc, caption)

    doc.add_heading('4. 2026-2030 operasyonel risk görünümü', level=1)
    nt = doc.add_table(rows=1, cols=6)
    style_table(nt)
    for i, h in enumerate(['Senaryo', 'Tepe %40 altı risk', 'Tepe ay', 'Tepe %30 altı risk', 'Tepe ay', 'En zayıf yaz P50']):
        nt.cell(0, i).text = h
    for row in nearterm.itertuples(index=False):
        r = nt.add_row().cells
        r[0].text = SCENARIO_LABELS.get(row.scenario, row.scenario)
        r[1].text = fmt_pct(row.max_monthly_prob_below_40_2026_2030_pct)
        r[2].text = str(row.peak_month_prob_below_40_date)
        r[3].text = fmt_pct(row.max_monthly_prob_below_30_2026_2030_pct)
        r[4].text = str(row.peak_month_prob_below_30_date)
        r[5].text = f"{int(row.lowest_p50_summer_year_2026_2030)} / {fmt_pct(row.lowest_p50_summer_fill_pct_2026_2030)}"

    add_bullet(doc, 'Temel senaryoda 2030 sonuna kadar %40 altında üç ay üst üste olasılığı %50 yi geçmiyor.')
    add_bullet(doc, 'Yönetim iyileşme senaryosu da yakın vadede benzer şekilde kalıcı düşük seviye üretmiyor.')
    hot = nearterm[nearterm['scenario'] == 'hot_dry_high_demand'].iloc[0]
    add_bullet(doc, f"Sıcak-kurak-yüksek talep senaryosunda %40 altı risk 2026-10 ayında %{hot['max_monthly_prob_below_40_2026_2030_pct']:.2f} düzeyine çıkıyor.")

    for path, caption, width in [
        (RISK_DIR / 'preferred_monthly_threshold_risk_2026_2030.png', 'Şekil 5. 2026-2030 döneminde aylık eşik altı olasılık yolları.', 6.1),
        (RISK_DIR / 'preferred_threshold_heatmap_40_2026_2030.png', 'Şekil 6. 2026-2030 aylık %40 altı risk ısı haritası.', 6.1),
    ]:
        doc.add_picture(str(path), width=Inches(width))
        add_caption(doc, caption)

    doc.add_heading('5. Teknik yorum ve sınırlar', level=1)
    add_bullet(doc, 'Referans evapotranspirasyon (ET0) bloğu hâlâ vekil radyasyon serileriyle çalışıyor. Aktinograf geldiğinde yeniden kurulmalı.')
    add_bullet(doc, 'Açık su yüzeyi buharlaşması, ET0 ile aynı şey gibi ele alınmıyor; modüler su bütçesi mantığı korunuyor.')
    add_bullet(doc, 'Kaynak bazlı yağış forcing denendi; fakat geleceğe taşınan uzaysal forcing zayıf olduğu için ana projeksiyon hattına alınmadı.')
    add_bullet(doc, 'Bugünkü en güvenilir kullanım şekli: yakın vadede rekonsiliasyonlu yol, orta-uzun vadede fizik-kısıtlı ensemble omurgası ve her ikisi için olasılıksal risk okuması.')

    d = doc.add_table(rows=1, cols=2)
    style_table(d, 'DCFCE7')
    d.cell(0,0).text = 'Senaryo'
    d.cell(0,1).text = 'Rekonsiliasyonun 2040 P50 etkisi (yp)'
    for row in compare[['scenario', 'delta_p50_endpoint_pp']].itertuples(index=False):
        r = d.add_row().cells
        r[0].text = SCENARIO_LABELS.get(row.scenario, row.scenario)
        r[1].text = fmt_num(row.delta_p50_endpoint_pp)

    doc.add_paragraph('Tablo yorumu: rekonsiliasyonun 2040 sonu etkisi çok küçük. Bu beklenen bir sonuçtur; düzeltme tasarımı gereği yalnızca yakın vadede etkili olup daha sonra sönümlenir.')

    doc.add_heading('6. Kaynakça', level=1)
    for row in refs.itertuples(index=False):
        doc.add_paragraph(f"[{row.source_id}] {row.title}. {row.organization_or_journal}. {row.year}. {row.url}")

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == '__main__':
    main()
