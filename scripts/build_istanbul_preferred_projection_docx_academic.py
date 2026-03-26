#!/usr/bin/env python3
from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path('/Users/yasinkaya/Hackhaton')
PREFERRED_DIR = ROOT / 'output' / 'istanbul_preferred_projection_2040'
RISK_DIR = PREFERRED_DIR / 'operational_risk_2030'
BENCH_PATH = ROOT / 'output' / 'istanbul_hybrid_physics_sourceaware_ensemble_2040' / 'ensemble_vs_benchmark_models.csv'
ITER_PATH = ROOT / 'output' / 'model_iteration_comparison_2026_03_12.csv'
BOTTOM_UP_SUMMARY_PATH = ROOT / 'output' / 'nearterm_bottom_up_reconciliation' / 'nearterm_bottom_up_reconciliation_summary.json'
SOURCE_REGISTRY = ROOT / 'research' / 'baraj_doluluk_hub' / 'registry' / 'sources' / 'external_sources.csv'
OUT_DOCX = ROOT / 'output' / 'doc' / 'istanbul_baraj_tercih_edilen_projeksiyon_raporu_akademik.docx'
USED_SOURCE_IDS = ['SRC-001', 'SRC-002', 'SRC-057', 'SRC-059', 'SRC-064', 'SRC-065', 'SRC-066', 'SRC-067', 'SRC-068', 'SRC-072', 'SRC-073']
SCENARIO_LABELS = {
    'wet_mild': 'Ilık-ıslak',
    'management_improvement': 'Yönetim iyileşme',
    'base': 'Temel',
    'hot_dry_high_demand': 'Sıcak-kurak-yüksek talep',
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def style_table(table, header_fill: str = 'D9EAF7') -> None:
    table.style = 'Table Grid'
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_fill)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style='List Bullet')
    p.add_run(text)


def add_caption(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def add_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)


def load_data() -> dict[str, object]:
    det = pd.read_csv(PREFERRED_DIR / 'deterministic_summary.csv')
    prob = pd.read_csv(PREFERRED_DIR / 'probabilistic_summary.csv')
    compare = pd.read_csv(PREFERRED_DIR / 'compare_with_baseprob.csv')
    nearterm = pd.read_csv(RISK_DIR / 'preferred_nearterm_risk_summary_2026_2030.csv')
    bench = pd.read_csv(BENCH_PATH)
    iterations = pd.read_csv(ITER_PATH)
    bottom_up = json.loads(BOTTOM_UP_SUMMARY_PATH.read_text(encoding='utf-8'))
    refs = pd.read_csv(SOURCE_REGISTRY)
    refs = refs[refs['source_id'].isin(USED_SOURCE_IDS)].copy()
    refs['sort_key'] = refs['source_id'].str.extract(r'(\d+)').astype(int)
    refs = refs.sort_values('sort_key')
    return {
        'det': det,
        'prob': prob,
        'compare': compare,
        'nearterm': nearterm,
        'bench': bench,
        'iterations': iterations,
        'bottom_up': bottom_up,
        'refs': refs,
    }


def fmt_pct(v, digits: int = 2) -> str:
    if v is None or v == '' or pd.isna(v):
        return '-'
    return f'%{float(v):.{digits}f}'


def fmt_num(v, digits: int = 2) -> str:
    if v is None or v == '' or pd.isna(v):
        return '-'
    return f'{float(v):.{digits}f}'


def main() -> None:
    data = load_data()
    det = data['det']
    prob = data['prob']
    compare = data['compare']
    nearterm = data['nearterm']
    bench = data['bench']
    iterations = data['iterations']
    bottom_up = data['bottom_up']
    refs = data['refs']

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('İstanbul Baraj Doluluğu İçin Tercih Edilen Projeksiyon Çerçevesi')
    r.bold = True
    r.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run('Yakın vade uzlaştırması ile güçlendirilmiş fizik-kısıtlı ve kaynak duyarlı hibrit yaklaşım')
    s.italic = True
    s.font.size = Pt(10.5)

    doc.add_heading('Özet', level=1)
    add_paragraph(doc, 'Bu belge, İstanbul toplam baraj doluluk oranı için geliştirilen projeksiyon sisteminde hangi model ailesinin tercih edildiğini, bu tercihin hangi test sonuçlarına dayandığını ve çıktının hangi koşullarda kullanılmasının uygun olduğunu özetlemektedir. Tercih edilen çözüm, yakın vadede baraj bazlı tahminlerle uzlaştırılmış bir yol ile orta-uzun vadede fizik-kısıtlı ve kaynak duyarlı hibrit ensemble omurgasını birlikte kullanmaktadır. Bu nedenle belge, hem teknik başarı ölçütlerini hem de kullanım sınırlarını birlikte raporlamaktadır.')

    chosen = bench[bench['model'] == 'hybrid_physics_ensemble_phys'].iloc[0]
    metrics = doc.add_table(rows=2, cols=3)
    style_table(metrics, '0F766E')
    headers = ['Tek adım hata', 'Çok adımlı ortalama hata', 'Fiziksel işaret testi']
    values = [fmt_num(chosen['one_step_rmse_pp']), fmt_num(chosen['mean_recursive_rmse_pp']), '4/4']
    for i in range(3):
        metrics.cell(0, i).text = headers[i]
        metrics.cell(1, i).text = values[i]
        set_cell_shading(metrics.cell(1, i), 'CCFBF1')

    doc.add_heading('1. Amaç ve yöntemsel tercih', level=1)
    add_paragraph(doc, 'Çalışmanın amacı, İstanbul toplam baraj doluluğu için yalnızca geçmiş seriyi ileri taşıyan bir tahmin üretmek değil; iklim, talep, işletme ve dış transfer baskılarını birlikte ele alabilen savunulabilir bir projeksiyon çerçevesi kurmaktır. Bu nedenle model seçimi yalnızca hata minimizasyonuna göre yapılmamış, aynı zamanda fiziksel işaret testleri ve yakın vade operasyonel tutarlılık da değerlendirmeye alınmıştır.')
    add_bullet(doc, 'Ana model ailesi: fizik-kısıtlı hibrit ensemble.')
    add_bullet(doc, 'Yakın vade düzeltmesi: kapasite ağırlıklı bottom-up baraj tahminleri ile uzlaştırma.')
    add_bullet(doc, 'Olasılıksal çıktı: medyan yol ile birlikte P10-P90 aralığı ve eşik altı risk tabloları.')

    doc.add_heading('2. Model seçimi ve elenen sürümler', level=1)
    add_paragraph(doc, 'Model karşılaştırmasında temel kriter, tek adımlı başarı ile çok adımlı kararlılık arasında dengeli bir çözüm üretmek olmuştur. Bu çerçevede `hybrid_physics_ensemble_phys`, hem düşük hata düzeyi hem de fiziksel işaret testlerinin tamamını geçmesi nedeniyle tercih edilmiştir.')

    table = doc.add_table(rows=1, cols=4)
    style_table(table)
    hdr = table.rows[0].cells
    hdr[0].text = 'Model'
    hdr[1].text = 'Tek adım hata'
    hdr[2].text = 'Çok adımlı hata'
    hdr[3].text = 'Fizik testi'
    for row in bench.sort_values('one_step_rmse_pp').head(6).itertuples(index=False):
        r = table.add_row().cells
        r[0].text = row.model
        r[1].text = fmt_num(row.one_step_rmse_pp)
        r[2].text = fmt_num(row.mean_recursive_rmse_pp)
        r[3].text = f"{int(row.physics_pass_count)}/4"

    v5 = iterations[iterations['model'] == 'ensemble_v5_sourcerain_phys'].iloc[0]
    add_paragraph(doc, f"Kaynak bazlı yağış forcing içeren `ensemble_v5_sourcerain_phys` kısa ufukta {v5['one_step_rmse_pp']:.2f} ile daha düşük hata vermiştir; ancak çok adımlı ortalama hata {v5['mean_recursive_rmse_pp']:.2f} düzeyine yükselmiştir. Bu nedenle söz konusu sürüm araştırma dalında tutulmuş, ana projeksiyon hattına alınmamıştır. Sonuç olarak mevcut tercih, kısa vadeli tek bir iyileşmeden ziyade daha dengeli bir genel performansa dayanır.")

    doc.add_heading('3. Yakın vade uzlaştırması', level=1)
    add_paragraph(doc, 'Top-down temel yol ile baraj bazlı bottom-up tahminler karşılaştırıldığında, yakın vadede sistematik bir iyimserlik görülmüştür. Bu fark doğrudan raporlanmakla kalınmamış, 2026-2029 dönemine uygulanan ve zamanla sönen bir düzeltme eğrisine dönüştürülmüştür. Böylece operasyonel kısa vade görünümü daha temkinli bir çizgiye çekilmiş, 2030 sonrası ise ana fizik-kısıtlı omurga korunmuştur.')
    add_bullet(doc, f"Bottom-up ile top-down temel yol arasındaki ortalama fark: {bottom_up['mean_gap_bottom_up_minus_topdown_pp']:.2f} yüzde puan.")
    add_bullet(doc, f"Yakın vadedeki en sert fark: {bottom_up['min_gap_bottom_up_minus_topdown_pp']:.2f} yüzde puan.")
    add_bullet(doc, 'Düzeltme 2030 sonrasında sıfıra yaklaşacak şekilde kurgulanmıştır.')

    for path, caption, width in [
        (ROOT / 'output' / 'nearterm_bottom_up_reconciliation' / 'nearterm_bottom_up_vs_topdown.png', 'Şekil 1. Baraj bazlı bottom-up tahminlerin kapasite ağırlıklı toplamı ile top-down temel yol arasındaki fark.', 6.0),
        (PREFERRED_DIR / 'reconciled_base_projection.png', 'Şekil 2. Uzlaştırılmış temel senaryo yolu. Düzeltme etkisi 2026-2029 döneminde gözlenmekte, sonrasında sönümlenmektedir.', 6.0),
    ]:
        doc.add_picture(str(path), width=Inches(width))
        add_caption(doc, caption)

    doc.add_heading('4. Deterministik ve olasılıksal senaryo çıktıları', level=1)
    merged = det.merge(prob, on='scenario', how='left')
    t = doc.add_table(rows=1, cols=6)
    style_table(t)
    for i, h in enumerate(['Senaryo', '2040 sonu', '2040 P10', '2040 P50', '2040 P90', 'İlk %40>=50 yılı']):
        t.cell(0, i).text = h
    for row in merged.itertuples(index=False):
        r = t.add_row().cells
        r[0].text = SCENARIO_LABELS.get(row.scenario, row.scenario)
        r[1].text = fmt_pct(row.end_fill_2040_12_pct)
        r[2].text = fmt_pct(row.p10_endpoint_2040_12_pct)
        r[3].text = fmt_pct(row.p50_endpoint_2040_12_pct)
        r[4].text = fmt_pct(row.p90_endpoint_2040_12_pct)
        r[5].text = '-' if pd.isna(row.first_year_prob_below_40_ge_50pct) else str(int(row.first_year_prob_below_40_ge_50pct))

    add_paragraph(doc, 'Deterministik sonuçlar ile olasılıksal bant birlikte okunduğunda, temel senaryoda 2040 medyan doluluk düzeyinin yaklaşık %34.69 civarında kaldığı, yönetim iyileşmesi senaryosunun ise iklim aynı kalsa bile sistemi yukarı taşıdığı görülmektedir. Buna karşılık sıcak-kurak-yüksek talep senaryosu, erken dönemde de alarm üreten tek agresif kırılma senaryosu olarak öne çıkmaktadır.')

    for path, caption, width in [
        (PREFERRED_DIR / 'reconciled_probabilistic_fan_paths_2026_2040.png', 'Şekil 3. Tercih edilen paketin olasılıksal senaryo yolları. Orta çizgi medyan yolu, bantlar belirsizlik aralığını göstermektedir.', 6.15),
        (PREFERRED_DIR / 'reconciled_vs_baseprob_base.png', 'Şekil 4. Rekonsiliasyon öncesi ve sonrası temel senaryo medyan yolu. Etki kısa vadede gözlenmekte, uzun vadede sönümlenmektedir.', 6.15),
    ]:
        doc.add_picture(str(path), width=Inches(width))
        add_caption(doc, caption)

    doc.add_heading('5. 2026-2030 operasyonel risk görünümü', level=1)
    nt = doc.add_table(rows=1, cols=6)
    style_table(nt)
    for i, h in enumerate(['Senaryo', 'Tepe %40 altı risk', 'Tepe ay', 'Tepe %30 altı risk', 'Tepe ay', 'En zayıf yaz P50']):
        nt.cell(0, i).text = h
    for row in nearterm.itertuples(index=False):
        r = nt.add_row().cells
        r[0].text = SCENARIO_LABELS.get(row.scenario, row.scenario)
        r[1].text = fmt_pct(row.max_monthly_prob_below_40_2026_2030_pct)
        r[2].text = str(row.peak_month_prob_below_40_date)
        r[3].text = fmt_pct(row.max_monthly_prob_below_30_2026_2030_pct)
        r[4].text = str(row.peak_month_prob_below_30_date)
        r[5].text = f"{int(row.lowest_p50_summer_year_2026_2030)} / {fmt_pct(row.lowest_p50_summer_fill_pct_2026_2030)}"

    add_paragraph(doc, 'Yakın vade risk analizinde temel ve yönetim iyileşme senaryoları için 2030 sonuna kadar %40 altında üç ay üst üste kalıcı stres olasılığı %50 düzeyini aşmamaktadır. Buna karşılık sıcak-kurak-yüksek talep senaryosu, 2026 yılından itibaren belirgin biçimde alarm üretmektedir. Bu durum, ana operasyonel baskının her senaryoda aynı olmadığını, özellikle talep ve iklim stresinin birlikte ele alınması gerektiğini göstermektedir.')
    for path, caption, width in [
        (RISK_DIR / 'preferred_monthly_threshold_risk_2026_2030.png', 'Şekil 5. 2026-2030 döneminde aylık eşik altı olasılık yolları.', 6.0),
        (RISK_DIR / 'preferred_threshold_heatmap_40_2026_2030.png', 'Şekil 6. 2026-2030 aylık %40 altı risk ısı haritası.', 6.0),
    ]:
        doc.add_picture(str(path), width=Inches(width))
        add_caption(doc, caption)

    doc.add_heading('6. Sınırlılıklar ve sonraki adımlar', level=1)
    add_bullet(doc, 'Referans evapotranspirasyon (ET0) bloğu hâlen vekil radyasyon serileriyle çalışmaktadır; aktinograf verisi geldiğinde yeniden kurulmalıdır.')
    add_bullet(doc, 'Açık su yüzeyi buharlaşması ET0 ile özdeş kabul edilmemekte, modüler su bütçesi mantığı korunmaktadır; ancak açık su yüzeyi alan eğrileri henüz eksiktir.')
    add_bullet(doc, 'Kaynak bazlı yağış forcing denenmiş, fakat geleceğe taşınan uzaysal forcing zayıf olduğunda model kararlılığı bozulmuştur.')
    add_bullet(doc, 'Bu nedenle bugünkü en uygun kullanım biçimi, yakın vadede uzlaştırılmış yol ve orta-uzun vadede fizik-kısıtlı ensemble omurgasının birlikte yorumlanmasıdır.')

    d = doc.add_table(rows=1, cols=2)
    style_table(d, 'DCFCE7')
    d.cell(0, 0).text = 'Senaryo'
    d.cell(0, 1).text = 'Rekonsiliasyonun 2040 P50 etkisi (yp)'
    for row in compare[['scenario', 'delta_p50_endpoint_pp']].itertuples(index=False):
        r = d.add_row().cells
        r[0].text = SCENARIO_LABELS.get(row.scenario, row.scenario)
        r[1].text = fmt_num(row.delta_p50_endpoint_pp)
    add_paragraph(doc, 'Tablo yorumu: rekonsiliasyonun 2040 sonu üzerindeki etkisi sınırlıdır. Bu sonuç beklenmektedir; çünkü söz konusu düzeltme tasarım gereği yakın vade operasyonel görünümü iyileştirmek için tanımlanmış ve zamanla sönümlenecek şekilde kurulmuştur.')

    doc.add_heading('Kaynakça', level=1)
    for row in refs.itertuples(index=False):
        doc.add_paragraph(f"[{row.source_id}] {row.title}. {row.organization_or_journal}. {row.year}. {row.url}")

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == '__main__':
    main()
